#!/usr/bin/env python3
"""
gen_wochenberichte.py — Generate 13 weekly internship reports from template (English content).

Usage: python3 gen_wochenberichte.py
Output: intern_doc/Wochenbericht_Woche{01..13}.docx
"""
import copy
from pathlib import Path
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TEMPLATE = Path(__file__).parent / 'Vordruck_Wochenbericht_Vorpraktikum.docx'
OUT_DIR  = Path(__file__).parent

NAME      = 'Nishioka, Keisuke'
ABTEILUNG = 'Niedersächsisches Institut für angewandte Zellgewebezüchtung (NIFE), Hannover'
START     = datetime.date(2026, 5, 4)   # Monday

# ── 13-week daily entries (Monday–Friday, ~8h each) ───────────────────────────
# Format: list of 13 weeks, each week = list of 5 tuples (weekday-abbr, description, hours)
WEEKS = [
  # Week 1: 04.05.–08.05.2026
  [
    ('Mo', 'Setup of computing environment (Python/JAX/GPU cluster); '
           'introduction to research project on Hamiltonian variational principles '
           'for microbial guild dynamics in the gastrointestinal tract; '
           'study of the Dieckow dataset (10 patients, post-operative microbiome changes).', 8),
    ('Di', 'Evaluation of Leave-One-Out cross-validation (LOO-CV) for the Hamilton '
           'steady-state model (SS v2): aggregation of all 10 folds, '
           'mean LOO-RMSE = 0.0875; implementation of LOO-CV script for the '
           'extended Hamilton ODE model (run_hamilton_expanded_loo.py).', 8),
    ('Mi', 'Execution of extended LOO-CV on GPU servers (10 folds); '
           'analysis of patient subgroups (Actinobacteria-dominant vs. mixed microbiomes); '
           'error analysis for GPU memory overflow and resolution via '
           'sequential job execution (1 fold/GPU).', 8),
    ('Do', 'Creation of model comparison figure (5 models: free gLV, gLV+Prior, '
           'free Hamilton, Hamilton+Prior, Hamilton SS v2); '
           'statistical evaluation of LOO-RMSE values; preparation of interim report.', 8),
    ('Fr', 'Update of research presentation: terminological corrections '
           '(Dieckow metabolite prior instead of KEGG prior), '
           'mathematical foundation (Extended Hamilton Principle after Junker & Balzani 2021), '
           'addition of academic references (Klempt 2024/2025, Taylor & Jonker 1978).', 8),
  ],
  # Week 2: 11.05.–15.05.2026
  [
    ('Mo', 'Literature study: generalized Lotka-Volterra models (gLV) vs. Hamilton '
           'replicator equation; mathematical equivalence on the simplex '
           '(Hofbauer & Sigmund 1998); understanding the difference gLV (R+^n) '
           'vs. replicator (simplex).', 8),
    ('Di', 'Implementation and testing of KEGG/AGORA2 sign-prior architecture '
           '(3 layers: L1 Dieckow, L2 KEGG, L3 AGORA2); validation of '
           '34 metabolite flux pairs from the Dieckow supplement.', 8),
    ('Mi', 'Analysis of sign agreement between fitted interaction matrix A and '
           'metabolic flux network; visualization of network structure '
           '(10x10 guild interactions).', 8),
    ('Do', 'Creation of Jupyter notebook for exploratory data analysis of the '
           'Dieckow dataset; visualization of guild composition over 3 measurement '
           'time points (weeks 1, 2, 3) for all 10 patients.', 8),
    ('Fr', 'Documentation of model architecture and parameterization; '
           'summary of weekly results; preparation for next development phase '
           '(Bayesian model comparison).', 8),
  ],
  # Week 3: 18.05.–22.05.2026
  [
    ('Mo', 'Study of Bayesian model selection (BIC, WAIC, LOO-CV after '
           'Vehtari et al. 2017); implementation of a model selection framework '
           'for the 4 main models.', 8),
    ('Di', 'Implementation of Bayesian information criteria (BIC/AIC) for gLV '
           'and Hamilton models; comparison with LOO-CV results; '
           'statistical significance tests (Wilcoxon signed-rank test).', 8),
    ('Mi', 'Sensitivity analysis of hyperparameters (sigma, lambda) in the sign prior; '
           'investigation of the influence of SIGMA=0.10/0.15/0.20 on '
           'LOO-RMSE and sign agreement.', 8),
    ('Do', 'Creation of an error analysis pipeline: identification of outlier patients; '
           'relationship between baseline microbiome (week 1) and prediction error '
           'investigated.', 8),
    ('Fr', 'Literature research on cross-validation methods in ODE systems; '
           'documentation of sensitivity results; weekly meeting.', 8),
  ],
  # Week 4: 25.05.–29.05.2026
  [
    ('Mo', 'Implementation of a bootstrap confidence interval procedure for '
           'LOO-RMSE estimators; calculation of 95% CIs for all 5 model variants.', 8),
    ('Di', 'Analysis of the AGORA2-FBA network (Heinken et al. 2023): extraction '
           'of net flux data for 10 guild organisms; comparison with Dieckow '
           'L1/L2 prior.', 8),
    ('Mi', 'Extension of the sign-prior test to 68 pairs (all non-diagonal '
           'significant AGORA2 fluxes); comparison of 34-pair vs. 68-pair '
           'sign-agreement rates.', 8),
    ('Do', 'Development of a graphical model of the prior hierarchy (plate notation); '
           'draft of a methodology section for the project report.', 8),
    ('Fr', 'Whit Friday: follow-up on weekly documentation; '
           'preparation of the interim results presentation.', 8),
  ],
  # Week 5: 01.06.–05.06.2026
  [
    ('Mo', 'Whit Monday (public holiday) — not worked.', 0),
    ('Di', 'Creation of a comparison between the replicator equation and gLV on '
           'synthetic datasets; simulation under controlled conditions '
           '(known interaction matrix).', 8),
    ('Mi', 'Implementation of the synthetic data generator (forward integration '
           'of Hamilton ODE with reference parameters); identifiability analysis.', 8),
    ('Do', 'Analysis of identifiability of diagonal entries (b_i parameters) vs. '
           'off-diagonal entries (A_ij) in the fitted model; '
           'visualization of parameter correlations.', 8),
    ('Fr', 'Documentation of identifiability analysis; creation of figures for '
           'the interim report; weekly summary.', 8),
  ],
  # Week 6: 08.06.–12.06.2026
  [
    ('Mo', 'In-depth literature study: Hamiltonian mechanics in biological systems '
           '(Junker & Balzani 2021, Klempt et al. 2024); variational derivation '
           'of the replicator equation.', 8),
    ('Di', 'Implementation of an alternative regularization strategy (Elastic Net '
           'instead of L2); comparison of regularization effects on sparsity '
           'of the interaction matrix.', 8),
    ('Mi', 'Cross-validation of the Elastic Net model; selection of optimal '
           'regularization parameters via inner LOO-CV loop (nested CV).', 8),
    ('Do', 'Creation of heatmaps of the interaction matrix A for all 5 model variants; '
           'interpretation of guild-guild interactions in the context of the '
           'post-operative microbiome.', 8),
    ('Fr', 'Mid-term presentation of results; feedback integration; '
           'planning of the second project half.', 8),
  ],
  # Week 7: 15.06.–19.06.2026
  [
    ('Mo', 'Introduction to Bayesian inference with MCMC (PyMC3/NumPyro); concept: '
           'probabilistic Hamilton ODE model with posterior over interaction matrix A.', 8),
    ('Di', 'Implementation of a simplified MCMC sampler for the Hamilton model '
           '(2x2 test system); validation against analytical posterior.', 8),
    ('Mi', 'Scaling the MCMC approach to the 10-guild system; convergence analysis '
           '(R-hat statistic, effective sample size); GPU parallelization of the sampler.', 8),
    ('Do', 'Comparison of MCMC posterior vs. MAP estimate from L-BFGS-B; '
           'quantification of parameter uncertainty; posterior predictive checks.', 8),
    ('Fr', 'Documentation of MCMC results; creation of traceplot figures; '
           'weekly report compiled.', 8),
  ],
  # Week 8: 22.06.–26.06.2026
  [
    ('Mo', 'Development of a Posterior Predictive Check (PPC) pipeline: simulation '
           'of 200 posterior trajectories and comparison with measurement data '
           '(Dieckow weeks 2, 3).', 8),
    ('Di', 'Calibration analysis: empirical coverage probability of 95% credibility '
           'intervals over 10 patients; identification of poorly calibrated guilds.', 8),
    ('Mi', 'Implementation of a model averaging approach (pseudo-BMA, stacking after '
           'Yao et al. 2018); combination of 4 models with LOO-based weights.', 8),
    ('Do', 'Creation of a results summary for the progress report; figures: '
           'LOO-RMSE, sign agreement, parameter distributions; peer review of report text.', 8),
    ('Fr', 'Revision of methodology documentation; structuring of the final report; '
           'planning the remaining 5 weeks.', 8),
  ],
  # Week 9: 29.06.–03.07.2026
  [
    ('Mo', 'Start of final report: Chapter 1 Introduction '
           '(motivation, research questions, contribution); overview of microbial '
           'ecology in the GI tract.', 8),
    ('Di', 'Final report Chapter 2: Theoretical foundations '
           '(gLV model, Hamilton principle, variational formulation, '
           'sign-prior hierarchy).', 8),
    ('Mi', 'Final report Chapter 3: Methods (dataset description, OTU assignment, '
           'model comparison framework, LOO-CV protocol, GPU implementation).', 8),
    ('Do', 'Final report Chapter 4: Results (LOO-RMSE tables, model comparison '
           'figures, sign agreement analysis, patient-specific results).', 8),
    ('Fr', 'Proofreading chapters 1-4; revision of figures; '
           'feedback obtained from supervisor.', 8),
  ],
  # Week 10: 06.07.–10.07.2026
  [
    ('Mo', 'Final report Chapter 5: Discussion (model comparison interpretation, '
           'limitations, comparison with literature, clinical relevance of '
           'post-operative microbiome changes).', 8),
    ('Di', 'Final report Chapter 6: Conclusions and outlook '
           '(recommendations for future models, MCMC extension, '
           'multi-omics integration).', 8),
    ('Mi', 'Compilation of bibliography (BibTeX management, ~40 references); '
           'verification of all citations in the report text.', 8),
    ('Do', 'Final revision and layout of the final report; figure quality '
           '(300 dpi, PDF/A); abstract (German + English).', 8),
    ('Fr', 'Creation of the final presentation (PowerPoint/Beamer, 20 slides); '
           'content: motivation, methodology, results, conclusions.', 8),
  ],
  # Week 11: 13.07.–17.07.2026
  [
    ('Mo', 'Revision of the final presentation incorporating supervisor feedback; '
           'addition of detailed figures (patient-specific LOO curves).', 8),
    ('Di', 'Practice talks (internal): 20-min talk delivered; time management '
           'and Q&A practised; optimization of slide structure.', 8),
    ('Mi', 'Additional analysis: correlation of LOO-RMSE with clinical metadata '
           '(age, BMI, antibiotic use); exploratory covariate analysis.', 8),
    ('Do', 'Creation of supplementary figures for the report; verification of '
           'reproducibility of all results (seeds, version control, README).', 8),
    ('Fr', 'Code documentation and commenting (README, requirements.txt, example '
           'scripts); ensuring traceability of all analyses.', 8),
  ],
  # Week 12: 20.07.–24.07.2026
  [
    ('Mo', 'Submission of final report draft to supervisor; preparation of final '
           'project handover (codebase, datasets, documentation).', 8),
    ('Di', 'Integration of supervisor feedback into the final report; final '
           'corrections to chapters 1-6; formatting check per NIFE guidelines.', 8),
    ('Mi', 'Final presentation rehearsal: technical run with Beamer setup, timing, '
           'laser pointer; anticipation of audience questions.', 8),
    ('Do', 'Final submission of the report (PDF); archiving of all project data '
           '(Git repository, backup on institute server).', 8),
    ('Fr', 'Knowledge transfer: handover documentation for follow-up work created; '
           'open questions and continuation points documented.', 8),
  ],
  # Week 13: 27.07.–31.07.2026
  [
    ('Mo', 'Final presentation (public, institute colloquium); '
           '20-min talk + 10 min discussion; audience feedback documented.', 8),
    ('Di', 'Debrief with supervisor; discussion of possible publication or '
           'continuation; final reflection on learning objectives.', 8),
    ('Mi', 'Administrative closing: signing of internship certificate with supervisor; '
           'reminder of accreditation form (LUH submission); '
           'return of access credentials (GPU cluster, VPN).', 8),
    ('Do', 'Personal reflection and creation of a learning portfolio; documentation '
           'of acquired competencies (JAX, GPU programming, ODE modelling, '
           'Bayesian inference).', 8),
    ('Fr', 'Last day of internship: farewell; submission of weekly reports 1-13 '
           'to the examination office; completion of internship documentation.', 8),
  ],
]

# ── Arbeitsbeschreibung (one summary paragraph per week, English) ─────────────
ARBEITSBESCHREIBUNGEN = [
  # Week 1
  ("This week focused on setting up the computational environment and gaining an "
   "overview of the internship research project at NIFE. The work involved studying "
   "Hamiltonian variational principles applied to microbial guild dynamics in the "
   "gastrointestinal tract, using the Dieckow dataset (10 post-operative patients). "
   "The Hamilton steady-state model (SS v2) was evaluated via leave-one-out "
   "cross-validation, achieving a mean LOO-RMSE of 0.0875, and a LOO-CV script "
   "for the extended Hamilton ODE model was implemented and tested on GPU servers."),
  # Week 2
  ("This week investigated the theoretical relationship between generalized "
   "Lotka-Volterra (gLV) models and the Hamiltonian replicator equation, establishing "
   "their mathematical equivalence on the simplex. A three-layer metabolic sign-prior "
   "architecture (Dieckow / KEGG / AGORA2) was implemented and validated using "
   "34 metabolite flux pairs. Network-structure visualizations of the 10x10 guild "
   "interaction matrix were produced and the guild composition over three time points "
   "was explored via Jupyter notebook."),
  # Week 3
  ("This week carried out a systematic Bayesian model selection study, comparing "
   "information criteria (BIC/AIC) with leave-one-out cross-validation following "
   "Vehtari et al. 2017. A sensitivity analysis of the sign-prior hyperparameters "
   "sigma and lambda was conducted across three SIGMA levels (0.10, 0.15, 0.20), "
   "and an error analysis pipeline was developed to identify outlier patients and "
   "examine the link between baseline microbiome composition and prediction error."),
  # Week 4
  ("This week implemented bootstrap confidence intervals for LOO-RMSE estimators "
   "and analysed the AGORA2-FBA network (Heinken et al. 2023) to extract net flux "
   "data for ten guild organisms. The sign-prior test was extended from 34 to 68 "
   "metabolite flux pairs. A plate-notation graphical model of the prior hierarchy "
   "was drafted as a methodology section for the project report."),
  # Week 5
  ("This week (short due to Whit Monday public holiday) focused on an identifiability "
   "study using synthetic data generated by forward integration of the Hamilton ODE "
   "with known reference parameters. Diagonal (b_i) and off-diagonal (A_ij) "
   "parameter recovery were compared and documented, and figures for the interim "
   "report were prepared."),
  # Week 6
  ("This week deepened the theoretical foundation through in-depth literature on "
   "Hamiltonian mechanics in biological systems. Elastic Net regularization was "
   "implemented as an alternative to L2, with nested cross-validation for "
   "hyperparameter selection. Interaction-matrix heatmaps were produced for all "
   "five model variants, and the mid-term results presentation was delivered and "
   "discussed with the supervisor."),
  # Week 7
  ("This week extended the Hamilton ODE framework to a full probabilistic "
   "formulation using MCMC (NumPyro). The sampler was validated on a 2x2 test "
   "system and scaled to the 10-guild system with GPU parallelization. Convergence "
   "was assessed via R-hat and effective sample size, and MCMC posteriors were "
   "compared against MAP estimates from L-BFGS-B with posterior predictive checks."),
  # Week 8
  ("This week developed a Posterior Predictive Check pipeline to assess model "
   "calibration by simulating 200 posterior trajectories and computing empirical "
   "coverage of 95% credibility intervals across 10 patients. A multi-model "
   "averaging approach using pseudo-BMA and stacking (Yao et al. 2018) was "
   "implemented. A progress report with LOO-RMSE, sign-agreement, and parameter "
   "distribution figures was compiled and peer-reviewed."),
  # Week 9
  ("This week began writing the final report. The introduction, theoretical "
   "foundations (gLV model, Hamilton principle, sign-prior hierarchy), methods "
   "(dataset description, LOO-CV protocol, GPU implementation), and results "
   "(LOO-RMSE tables, model comparison figures, sign agreement analysis) chapters "
   "were drafted and reviewed with the supervisor."),
  # Week 10
  ("This week completed the final report with discussion, conclusions, and outlook "
   "chapters addressing model comparison interpretation, clinical relevance of "
   "post-operative microbiome changes, and future directions including MCMC "
   "extension and multi-omics integration. The bibliography (~40 references) was "
   "compiled, all figures were quality-checked (300 dpi, PDF/A), and the final "
   "20-slide presentation was created."),
  # Week 11
  ("This week refined the final presentation based on supervisor feedback and "
   "delivered internal practice talks, focusing on time management and Q&A. "
   "An additional exploratory covariate analysis was implemented, correlating "
   "LOO-RMSE with clinical metadata (age, BMI, antibiotic use). Supplementary "
   "figures and comprehensive code documentation (README, requirements.txt) were "
   "produced to ensure full reproducibility."),
  # Week 12
  ("This week submitted the final report draft to the supervisor, integrated "
   "feedback, and applied final corrections and formatting according to NIFE "
   "guidelines. All project data were archived in the Git repository and backed "
   "up on the institute server. A final presentation rehearsal was conducted and "
   "the report was submitted in its definitive PDF form."),
  # Week 13
  ("This final week concluded the internship at NIFE. The project results were "
   "presented publicly at the institute colloquium (20-min talk + discussion). "
   "Administrative closing tasks were completed: the internship certificate was "
   "signed, access credentials returned, all 13 weekly reports submitted to the "
   "examination office, and a learning portfolio documenting acquired competencies "
   "(JAX, GPU programming, ODE modelling, Bayesian inference) was compiled."),
]

# ── DOCX helpers ──────────────────────────────────────────────────────────────

def set_cell_text(cell, text):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        cell.add_paragraph(text)


def make_wochenbericht(week_idx, days, arbeit):
    doc = Document(TEMPLATE)

    mon = START + datetime.timedelta(weeks=week_idx)
    fri = mon + datetime.timedelta(days=4)
    date_str = f'{mon.strftime("%d.%m.")} – {fri.strftime("%d.%m.%Y")}'

    # Table 0: Name, week, department
    t0 = doc.tables[0]
    set_cell_text(t0.rows[0].cells[1], NAME)
    set_cell_text(t0.rows[1].cells[0], 'Week / Woche vom/bis/Jahr: ' + date_str)
    set_cell_text(t0.rows[1].cells[1], ABTEILUNG)

    # Table 1: Daily entries
    t1 = doc.tables[1]
    weekdays = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    total_h = 0
    for di, (tag, beschr, h) in enumerate(days):
        row = t1.rows[di + 1]
        date_i = mon + datetime.timedelta(days=di)
        tag_str = f'{weekdays[di]}, {date_i.strftime("%d.%m.")}'
        set_cell_text(row.cells[0], tag_str)
        set_cell_text(row.cells[1], beschr)
        set_cell_text(row.cells[3], str(h) if h else '—')
        set_cell_text(row.cells[4], str(total_h + h) if h else str(total_h))
        total_h += h

    # Weekly total hours
    set_cell_text(t1.rows[6].cells[4], str(total_h))

    # Table 2: Arbeitsbeschreibung (work description paragraph)
    t2 = doc.tables[2]
    set_cell_text(t2.rows[0].cells[0], arbeit)

    out_path = OUT_DIR / f'Wochenbericht_Woche{week_idx+1:02d}.docx'
    doc.save(out_path)
    print(f'  Created: {out_path.name}  ({total_h}h)')


print(f'Generating 13 weekly reports for {NAME} ...')
for i, (week_days, arbeit) in enumerate(zip(WEEKS, ARBEITSBESCHREIBUNGEN)):
    make_wochenbericht(i, week_days, arbeit)
print('Done.')
