#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import re
from dataclasses import dataclass
from pathlib import Path

import pandas as pd


_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _sanitize_text(s: str) -> str:
    s = str(s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = _CONTROL_CHARS_RE.sub("", s)
    return s.strip()


def _read_pdf_text(path: Path) -> str:
    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            t = _sanitize_text(t)
            if t:
                chunks.append(t)
    return "\n".join(chunks)


def _read_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(_sanitize_text(p.text))
    for table in doc.tables:
        for row in table.rows:
            cells = [_sanitize_text(c.text) for c in row.cells]
            line = "\t".join([c for c in cells if c])
            if line:
                parts.append(line)
    return "\n".join(parts)


@dataclass(frozen=True)
class ExtractedField:
    field: str
    value: str
    source: str
    evidence: str


def _extract_fields(text: str, source: str) -> tuple[list[ExtractedField], list[ExtractedField]]:
    found: list[ExtractedField] = []
    missing: list[ExtractedField] = []

    patterns: list[tuple[str, re.Pattern[str]]] = [
        ("material", re.compile(r"\b(Ti6Al4V|Y[- ]?TZP|titanium)\b", re.IGNORECASE)),
        ("timepoints_days", re.compile(r"\b(\d+)\s*d\b", re.IGNORECASE)),
        ("device_oral_splint", re.compile(r"\boral\s+splint\b", re.IGNORECASE)),
        ("device_24_well_plate", re.compile(r"\b24\s*well\s*plate\b", re.IGNORECASE)),
        ("medium_BHI", re.compile(r"\bBHI\b", re.IGNORECASE)),
        ("buffer_PIPES", re.compile(r"\bPIPES\b", re.IGNORECASE)),
        ("sucrose", re.compile(r"\bsucrose\b", re.IGNORECASE)),
        ("sequencing_platform", re.compile(r"\bIllumina\b|\bMiSeq\b|\bNovaSeq\b", re.IGNORECASE)),
        ("amplicon_region", re.compile(r"\bV[1-9]\s*[-–]\s*V[1-9]\b", re.IGNORECASE)),
        ("dada2", re.compile(r"\bDADA2\b", re.IGNORECASE)),
        ("qiime2", re.compile(r"\bQIIME\s*2\b|\bQIIME2\b", re.IGNORECASE)),
    ]

    for field, pat in patterns:
        m = pat.findall(text)
        if not m:
            missing.append(ExtractedField(field=field, value="not_found", source=source, evidence="not_found"))
            continue
        vals: list[str] = []
        if isinstance(m[0], tuple):
            for tup in m:
                for x in tup:
                    if x:
                        vals.append(str(x))
        else:
            vals = [str(x) for x in m if str(x)]
        norm_vals = []
        for v in vals:
            v2 = v.strip()
            if field == "material":
                if re.fullmatch(r"y[- ]?tzp", v2, flags=re.IGNORECASE):
                    v2 = "Y-TZP"
                elif re.fullmatch(r"ti6al4v", v2, flags=re.IGNORECASE):
                    v2 = "Ti6Al4V"
                elif re.fullmatch(r"titanium", v2, flags=re.IGNORECASE):
                    v2 = "titanium"
            norm_vals.append(v2)
        unique = []
        for v in norm_vals:
            if v not in unique:
                unique.append(v)
        evidence = pat.pattern
        found.append(ExtractedField(field=field, value=";".join(unique), source=source, evidence=evidence))
    return found, missing


def write_kv_tsv(path: Path, rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="\n", encoding="utf-8") as f:
        w = csv.writer(f, delimiter="\t", lineterminator="\n")
        w.writerow(["sample_alias", "field", "value", "source_file", "evidence"])
        for r in rows:
            w.writerow([r["sample_alias"], r["field"], r["value"], r["source_file"], r["evidence"]])


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--pdf", type=Path, default=Path(__file__).parent / "docs" / "ZJOM_16_2424227.pdf")
    ap.add_argument("--docx", type=Path, default=Path(__file__).parent / "docs" / "ZJOM_A_2424227_SM9092.docx")
    ap.add_argument("--combined-meta", type=Path, default=Path(__file__).parent / "data" / "combined_meta.tsv")
    ap.add_argument("--out", type=Path, default=Path(__file__).parent / "data" / "structured_supp.tsv")
    ap.add_argument("--manual-review", type=Path, default=Path(__file__).parent / "data" / "manual_review.tsv")
    args = ap.parse_args()

    meta = pd.read_csv(args.combined_meta, sep="\t")
    if "sample_alias" not in meta.columns:
        raise ValueError("combined_meta.tsv missing sample_alias")
    prjna_aliases = sorted({a for a in meta["sample_alias"].astype(str).tolist() if a.startswith("NU_")})
    if not prjna_aliases:
        raise ValueError("No PRJNA1159109 sample_alias (NU_*) found in combined_meta.tsv")

    pdf_text = _read_pdf_text(args.pdf)
    docx_text = _read_docx_text(args.docx)

    found_pdf, missing_pdf = _extract_fields(pdf_text, source=str(args.pdf))
    found_docx, missing_docx = _extract_fields(docx_text, source=str(args.docx))

    merged_fields: dict[str, ExtractedField] = {}
    for f in found_pdf + found_docx:
        if f.field not in merged_fields:
            merged_fields[f.field] = f
        else:
            prev = merged_fields[f.field]
            if prev.value == f.value:
                continue
            parts = []
            for v in [prev.value, f.value]:
                for x in str(v).split(";"):
                    x2 = _sanitize_text(x)
                    if x2:
                        parts.append(x2)
            uniq = []
            for x in parts:
                if x not in uniq:
                    uniq.append(x)
            merged_fields[f.field] = ExtractedField(
                field=f.field,
                value=";".join(uniq),
                source=prev.source + ";" + f.source,
                evidence=prev.evidence + ";" + f.evidence,
            )

    structured_rows: list[dict[str, str]] = []
    for a in prjna_aliases:
        for f in merged_fields.values():
            structured_rows.append(
                {
                    "sample_alias": a,
                    "field": f.field,
                    "value": _sanitize_text(f.value),
                    "source_file": _sanitize_text(f.source),
                    "evidence": _sanitize_text(f.evidence),
                }
            )

    manual_rows: list[dict[str, str]] = []
    for m in missing_pdf + missing_docx:
        manual_rows.append(
            {
                "sample_alias": "not_provided",
                "field": m.field,
                "value": "manual_check_required",
                "source_file": _sanitize_text(m.source),
                "evidence": _sanitize_text(m.evidence),
            }
        )

    write_kv_tsv(args.out, structured_rows)
    write_kv_tsv(args.manual_review, manual_rows)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
